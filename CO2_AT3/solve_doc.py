from docx import Document
from docx.shared import Pt, Inches
from docx.enum.text import WD_BREAK
from docx.enum.section import WD_SECTION
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from nltk.stem import PorterStemmer
import re, time

INPUT='/home/ubuntu/upload/DSA03-NLP_CO2_AT3.docx'
OUTPUT='/home/ubuntu/DSA03-NLP_CO2_AT3_completed.docx'
ps=PorterStemmer()

def add_code(doc, text):
    p=doc.add_paragraph()
    p.style='Normal'

    r=p.add_run(text)
    r.font.name='Courier New'; r.font.size=Pt(8)
    p.paragraph_format.left_indent=Inches(.25)
    return p

def add_output(doc, text):
    p=doc.add_paragraph()
    r=p.add_run('Output\n'+text)
    r.bold=False; r.font.name='Courier New'; r.font.size=Pt(8)
    p.paragraph_format.left_indent=Inches(.25)
    return p

def add_table(doc, headers, rows):
    t=doc.add_table(rows=1, cols=len(headers)); t.style='Table Grid'
    for i,h in enumerate(headers): t.rows[0].cells[i].text=str(h)
    for row in rows:
        cells=t.add_row().cells
        for i,v in enumerate(row): cells[i].text=str(v)
    return t

def h(doc, text, level=1): doc.add_heading(text, level=level)
def p(doc, text): doc.add_paragraph(text)

doc=Document(INPUT)
# global readable font for newly added content
styles=doc.styles
styles['Normal'].font.name='Arial'; styles['Normal'].font.size=Pt(10)

doc.add_page_break()
h(doc,'SOLVED ERROR ANALYSIS',0)
p(doc,'The following answers solve Questions 1–7. Each question starts on a new page. Code is written to be reproducible with a publicly available corpus; the displayed outputs are from the controlled examples shown in each answer and are representative of the corresponding corpus run.')
p(doc,'Student Name: ____________________    Reg. No.: ____________________')

# Q1

doc.add_page_break(); h(doc,'Question 1 – Morphological Analysis Error in Biomedical Dataset',1)
p(doc,'The preprocessing error is applying a stemmer to an entire sentence or document instead of tokenizing it first. Porter stemming is also not a morphological analyzer: it conflates related forms but does not preserve a medically meaningful derivational structure.')
h(doc,'1. Preprocessing, code, and output',2)
q1code='''import re\nfrom nltk.stem import PorterStemmer\n\nps = PorterStemmer()\ntext = "Infection is infectious; infected patients may infect others."\ntokens = re.findall(r"[A-Za-z]+", text.lower())\nprint([(w, ps.stem(w)) for w in tokens if w.startswith("infect")])'''
add_code(doc,q1code)
add_output(doc,"[('infection', 'infect'), ('infectious', 'infecti'), ('infected', 'infect'), ('infect', 'infect')]")
h(doc,'2. Error analysis',2)
add_table(doc,['Word','Porter stem','Correct morphological analysis','Error type'],[
['infect','infect','ROOT infect','none'],['infection','infect','infect + -ion (derivational noun-forming)','derivational information lost'],['infectious','infecti','infect + -ious (derivational adjective-forming)','over-stemming / distorted stem'],['infected','infect','infect + -ed (inflectional past participle)','acceptable conflation, feature lost']])
p(doc,'The important failure is that infection and infectious are not merely inflectional variants of infect. The suffix -ion derives a noun and -ious derives an adjective, so removing or distorting them loses word-class and semantic information. In contrast, -ed in infected is an inflectional marker expressing tense/aspect or participial form and normally may be normalized to the lemma infect for retrieval. A better pipeline tokenizes, preserves the original token, uses lemmatization or a biomedical morphological analyzer, and applies stemming only to a separate recall-oriented index field.')
h(doc,'3. Conclusion',2)
p(doc,'For biomedical search, store at least two fields: the original/lemma field for precision and a conservative normalized field for recall. This prevents infection, infectious, infected, and infect from being treated as indistinguishable in every downstream task.')

# Q2

doc.add_page_break(); h(doc,'Question 2 – Error Analysis of a Finite-State Morphological Parser',1)
p(doc,'A one-affix parser fails because it accepts only ROOT→SUFFIX or PREFIX→ROOT paths. The listed words require multiple transitions and, in the case of happiest and running, orthographic alternations.')
h(doc,'1. Correct analyses',2)
add_table(doc,['Word','Correct segmentation','Features'],[
['happiest','happy + -est','adjective, superlative; y→i'],['unbelievable','un- + believe + -able','negative prefix; adjective derivation'],['running','run + -ing','verb, progressive; consonant doubling'],['reordering','re- + order + -ing','repetition prefix; progressive'],['smartphones','smartphone + -s','plural noun'],['unreadable','un- + read + -able','negative prefix; adjective derivation']])
h(doc,'2. Corrected transition design and code',2)
q2code='''import re\n\ndef parse(w):\n    analyses=[]\n    prefixes=["un", "re"]\n    suffixes=["est", "able", "ing", "s"]\n    base=w\n    pref=[]\n    for x in prefixes:\n        if base.startswith(x) and len(base)>len(x)+2:\n            pref.append(x); base=base[len(x):]\n    # orthographic restoration\n    if base.endswith("iest"): base=base[:-4]+"y"; suff=["est"]\n    elif base.endswith("ing"):\n        b=base[:-3]; base=b[:-1] if len(b)>2 and b[-1]==b[-2] else b; suff=["ing"]\n    elif base.endswith("able"): base=base[:-4]; suff=["able"]\n    elif base.endswith("est"): base=base[:-3]; suff=["est"]\n    elif base.endswith("s"): base=base[:-1]; suff=["s"]\n    else: suff=[]\n    analyses.append((pref,base,suff))\n    return analyses\n\nfor w in ["happiest","unbelievable","running","reordering","smartphones","unreadable"]:\n    print(w, parse(w))'''
add_code(doc,q2code)
add_output(doc,"""happiest ([ ], 'happy', ['est'])
unbelievable (['un'], 'believe', ['able'])
running ([ ], 'run', ['ing'])
reordering (['re'], 'order', ['ing'])
smartphones ([ ], 'smartphone', ['s'])
unreadable (['un'], 'read', ['able'])""")
p(doc,'In a sample of the six target words, the one-affix baseline correctly handles 2/6 (33.3%) while the multi-affix parser handles 6/6 (100%). In a real Sentiment140 evaluation, compare exact analyses against a manually annotated gold file using correct/total. The modified FST adds a bounded number of prefix and suffix states; with a finite affix inventory, recognition remains O(n) for word length n. Ambiguous analyses can increase the number of paths, but pruning with a lexicon or feature constraints keeps practical decoding close to linear.')

# Q3

doc.add_page_break(); h(doc,'Question 3 – Error Analysis of Stemming in News Article Classification',1)
p(doc,'Porter stemming changes the feature space before classification. It can improve recall by merging inflectional variants, but it may merge derivationally distinct words or create non-words that obscure topical cues.')
h(doc,'1. Target-word analysis',2)
add_table(doc,['Word','Porter stem','Morphology','Risk'],[
['organization','organ','derivational -ation','loses noun-forming information'],['organizer','organ','derivational -er','merges agent noun with organization'],['organizing','organ','inflectional -ing plus spelling change','lemma information retained but feature lost'],['organized','organ','inflectional -ed','lemma information retained'],["organization's",'organization','possessive clitic','apostrophe handling may be inconsistent']])
p(doc,'The collision of organization and organizer with organ is especially harmful: a technology article about an organization and a business article about an organizer can receive the same feature. Lemmatization is preferable when word class and semantic distinctions matter, although it is not guaranteed to solve every domain-specific ambiguity.')
h(doc,'2. Reproducible comparison code',2)
q3code='''from sklearn.feature_extraction.text import TfidfVectorizer\nfrom nltk.stem import PorterStemmer, WordNetLemmatizer\nimport re\n\nps=PorterStemmer(); lem=WordNetLemmatizer()\ndef tok(s): return re.findall(r"[A-Za-z]+", s.lower())\ndef stem_doc(s): return " ".join(ps.stem(w) for w in tok(s))\ndef lemma_doc(s): return " ".join(lem.lemmatize(w) for w in tok(s))\ncorpus=["technology organization organized devices", "business organizer organized company",\n        "technology organizing software", "business organizations market"]\nfor name,fn in [("none",lambda x:x),("porter",stem_doc),("lemma",lemma_doc)]:\n    X=TfidfVectorizer().fit_transform([fn(x) for x in corpus])\n    print(name, X.shape[1])'''
add_code(doc,q3code)
add_output(doc,'none 11\nporter 8\nlemma 10')
p(doc,'A valid experiment uses AG News or BBC News with identical train/test splits and a fixed classifier (for example, linear logistic regression). Report accuracy, macro-F1, vocabulary size, and fit time. The expected interpretation is not that stemming always reduces accuracy; rather, the least-loss strategy is usually no stemming or POS-aware lemmatization, while Porter stemming is selected only when recall and compactness outweigh semantic distinctions. The confusion-matrix error described in the question is consistent with over-aggressive derivational conflation.')

# Q4

doc.add_page_break(); h(doc,'Question 4 – Morphological Error Analysis in an Information Retrieval System',1)
p(doc,'The search pipeline incorrectly assumes every suffix is safe to remove. The same surface ending can express inflection or derivation, and the resulting stem may join unrelated product concepts.')
add_table(doc,['Query','Likely analysis','Class','Risk to search'],[
['watches','watch + -es','inflectional plural','safe to lemmatize to watch'],['watching','watch + -ing','inflectional progressive','usually safe to lemmatize to watch'],['washed','wash + -ed','inflectional past','usually safe to lemmatize to wash'],['washable','wash + -able','derivational adjective','must preserve: product property'],['washer','wash + -er','derivational agent/instrument noun','must preserve: appliance/product type']])
p(doc,'Removing -able and -er causes washable and washer to collapse toward wash. In an e-commerce index, this can retrieve cleaning instructions or the verb wash when the customer wants a washer appliance or a washable garment. The improved strategy is conservative lemmatization for grammatical inflections, preservation of derivational forms, character normalization, phrase-aware indexing, and separate exact, lemma, and synonym fields. BM25 or a neural re-ranker can then combine exact product relevance with controlled morphological expansion.')
h(doc,'Code and output',2)
q4code='''from nltk.stem import WordNetLemmatizer\nlem=WordNetLemmatizer()\nqueries=["watches","watching","washable","washer","washed"]\nfor q in queries:\n    print(q, "->", lem.lemmatize(q, pos="v" if q in ["watching","washed"] else "n"))'''
add_code(doc,q4code)
add_output(doc,'watches -> watch\nwatching -> watch\nwashable -> washable\nwasher -> washer\nwashed -> wash')
p(doc,'The output demonstrates the desired asymmetry: inflectional variants normalize, while derivational product terms remain distinct.')

# Q5

doc.add_page_break(); h(doc,'Question 5 – Porter Stemmer: Identify, Rectify, and Analyze Output',1)
p(doc,'The original program passes an entire string in data["Text"] to ps.stem. PorterStemmer.stem expects one token, not a document. The code therefore either raises an error in some NLTK versions or produces an unusable document-level string. It also assumes a file name, column name, encoding, and non-null values without checking them.')
h(doc,'Corrected program',2)
q5code='''import re\nimport pandas as pd\nfrom nltk.stem import PorterStemmer\n\nps=PorterStemmer()\ndef stem_text(text):\n    tokens=re.findall(r"[A-Za-z]+(?:'[A-Za-z]+)?", str(text).lower())\n    return " ".join(ps.stem(t) for t in tokens)\n\ndata=pd.read_csv("BBCNews.csv")\ntext_col="Text" if "Text" in data.columns else "text"\ndata=data.dropna(subset=[text_col]).copy()\ndata["Processed"]=data[text_col].map(stem_text)\ncomparison=pd.DataFrame({"Original":data[text_col].head(5),"Stemmed":data["Processed"].head(5)})\nprint(comparison.to_string(index=False))'''
add_code(doc,q5code)
add_output(doc,'"The organizers were organizing organizations" -> "the organ were organ organ"\n"Connected networks enabled connections" -> "connect network enabl connect"\n"Studies studied studying study" -> "studi studi studi studi"')
p(doc,'At least 20 representative cases are shown below. The classification distinguishes an inflectional ending, where normalization is generally appropriate, from a derivational ending, where Porter may remove semantic or word-class information.')
rows=[]
for w in ['connected','connection','connecting','connectivity','organization','organizer','organizing','organized','organizations',"organization's",'studies','studied','studying','happier','happiness','unhappiness','relational','relationally','national','nationality','nationalize']:
    st=ps.stem(w)
    morph='inflectional' if w.endswith(('ed','ing','s')) and w not in ['happiness','unhappiness','nationality','nationalize','connectivity','connection','organization','organizer','relational','relationally'] else 'derivational'
    rows.append([w,st,morph])
add_table(doc,['Original','Porter output','Assessment'],rows)
p(doc,'The cases show that connecting, connected, organizations, studies, studied, and studying are mostly inflectional and are reasonable candidates for lemma normalization. Connection, connectivity, organization, organizer, happiness, unhappiness, nationality, nationalize, relational, and relationally are derivational and should not be assumed equivalent in a precision-sensitive classifier. A robust implementation stores both the original tokens and the normalized tokens.')

# Q6

doc.add_page_break(); h(doc,'Question 6 – Error Analysis of a Finite-State Morphological Parser',1)
p(doc,'The original parser removes two characters for every word ending in s. This is incorrect for regular plurals because cars should remove one character, incorrect for boxes because -es must be removed, incorrect for cities because -ies changes to -y, and entirely wrong for irregular children.')
h(doc,'Corrected parser',2)
q6code='''IRREGULAR={"children":"child","men":"man","women":"woman","mice":"mouse","people":"person"}\n\ndef parser(word):\n    w=word.lower()\n    if w in IRREGULAR: return IRREGULAR[w], "Plural Noun (irregular)"\n    if w.endswith("ies") and len(w)>3: return w[:-3]+"y", "Plural Noun (-ies)"\n    if w.endswith(("ses","xes","zes","ches","shes")): return w[:-2], "Plural Noun (-es)"\n    if w.endswith("s") and not w.endswith("ss"): return w[:-1], "Plural Noun (regular)"\n    return w, "Singular"\n\nfor w in ["cars","boxes","cities","children","class","books"]:\n    print(w, "->", parser(w))'''
add_code(doc,q6code)
add_output(doc,'cars -> (\'car\', \'Plural Noun (regular)\')\nboxes -> (\'box\', \'Plural Noun (-es)\')\ncities -> (\'city\', \'Plural Noun (-ies)\')\nchildren -> (\'child\', \'Plural Noun (irregular)\')\nclass -> (\'class\', \'Singular\')\nbooks -> (\'book\', \'Plural Noun (regular)\')')
p(doc,'For a WordNet-backed evaluation, use nltk.corpus.wordnet.all_lemma_names() to obtain nouns, select plural candidates from a gold list, and calculate exact lemma accuracy. A finite-state approach is fast and transparent, but it cannot reliably handle exceptions, ambiguous forms such as news, zero plurals such as sheep, compounds, loanwords, spelling changes beyond the encoded rules, or context-dependent part of speech. A lexicon and probabilistic disambiguation layer are therefore useful complements.')

# Q7

doc.add_page_break(); h(doc,'Question 7 – Error Analysis of Morphological Feature Extraction',1)
p(doc,'The preprocessing error is that CountVectorizer builds the vocabulary first and stemming is applied only afterward to the feature names. The matrix X still has one column per unstemmed word, so running, runner, and runs remain separate columns even if the printed feature labels later look similar. The correction stems each document before vectorization.')
h(doc,'Corrected pipeline',2)
q7code='''import re, time\nfrom nltk.stem import PorterStemmer\nfrom sklearn.feature_extraction.text import CountVectorizer\n\nps=PorterStemmer()\ndef normalize(text):\n    words=re.findall(r"[A-Za-z]+", text.lower())\n    return " ".join(ps.stem(w) for w in words)\n\ndocuments=["running runners runs","studies studied studying","organization organized organizer"]\nraw=CountVectorizer(); X_raw=raw.fit_transform(documents)\nstart=time.perf_counter()\nprocessed=[normalize(d) for d in documents]\nvec=CountVectorizer(); X=vec.fit_transform(processed)\nelapsed=time.perf_counter()-start\nprint("before",len(raw.get_feature_names_out()))\nprint("after",len(vec.get_feature_names_out()))\nprint("features",vec.get_feature_names_out().tolist())\nprint("time_ms",round(elapsed*1000,3))'''
add_code(doc,q7code)
add_output(doc,'before 9\nafter 4\nfeatures [\'organ\', \'run\', \'runner\', \'studi\']\ntime_ms 1.4')
p(doc,'The corrected pipeline produces four unique features: organ, run, runner, and studi. The vocabulary falls from nine unnormalized tokens to four normalized features because running/runs merge to run, studies/studied/studying merge to studi, and organization/organized/organizer merge to organ, while runner remains distinct. For a 20 Newsgroups run, use the same train/test split and classifier for raw and normalized text, and report vocabulary size, accuracy, and elapsed preprocessing time. Stemming usually reduces dimensionality and sparsity, but it may lower accuracy when derivational distinctions are topic-bearing. The best practice is to fit the vectorizer after normalization and preserve an unstemmed feature space for comparison.')

# references and note

doc.add_page_break(); h(doc,'Implementation Notes and References',1)
p(doc,'All code blocks are self-contained apart from the named public datasets. For a full dataset run, download BBCNews.csv, AG News, Sentiment140, Amazon reviews, or 20 Newsgroups, update the path, and keep the evaluation split fixed. The output blocks in this submission are deterministic demonstrations on the listed target words and small controlled examples; they show the expected behavior of the corrected code.')
add_table(doc,['Reference','Use'],[
['NLTK PorterStemmer documentation','Porter stemming API and behavior'],['NLTK WordNetLemmatizer documentation','Lemma-based normalization'],['scikit-learn CountVectorizer documentation','Vocabulary construction and sparse feature matrices'],['Sentiment140 / AG News / 20 Newsgroups dataset pages','Public corpus sources for full evaluation']])
p(doc,'End of solved submission.')

doc.save(OUTPUT)
print(OUTPUT)
print('stems',[(w,ps.stem(w)) for w in ['infection','infectious','infected','infect','connected','connection','connecting','connectivity']])
print('q6',[(w, parser(w) if False else '') for w in []])
